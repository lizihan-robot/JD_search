import sys,os
import docx 
from win32com.client import Dispatch
from win32com.client import DispatchEx
import pythoncom
import win32com

from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfpage import PDFTextExtractionNotAllowed
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.pdfdevice import PDFDevice
from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator
from threading import Thread

class load_file():
    def get_list_from_lib(self,lib_path, key_world_list):
        result_list=[0]
        find_key = True
        try:
            JD_list = os.listdir(lib_path)
        except:
            print("路径出错")
        
        for JD in JD_list: #循环遍历目录
            JD_path_name = os.path.join(lib_path,JD)
            print("----------------------------JD_path_name:",JD_path_name)
            if os.path.isdir(JD_path_name):
                print("目录中包含文件夹",JD_path_name)
                return EnvironmentError
            JD_msg = self.get_JD_msg(JD_path_name)
            result_list[0]+=1
            for key_world in key_world_list: #循环遍历关键词
                if JD_msg.find(key_world)!=-1:
                    find_key=True
                    print("已找到",key_world)
                else:
                    find_key = False
                    print("未找到",key_world)
                    break
            if find_key:
                result_list.append(JD_path_name)
                print("YYYYYYY在 %s 简历中搜索到关键词 %s"%(JD,key_world_list))
            else:
                print("NNNNNNN在 %s 简历中未找到关键词 %s"%(JD,key_world_list))
        print('result_list:',result_list)
        print("已查找了%s个文件"%result_list[0])
        return result_list

    def get_JD_msg(self,JD_path_name):
        print("JD_path_name:",os.path.splitext(JD_path_name)[1])
        if os.path.splitext(JD_path_name)[1] == ".docx":
            return self.get_text_docx(JD_path_name)
        if os.path.splitext(JD_path_name)[1] == ".pdf":
            return self.get_text_pdf(JD_path_name)

    def get_text_pdf(self,JD_path_name):
        "获取pdf内所有字符串"
        para_text =""
        fp=open(JD_path_name,"rb")
        parser=PDFParser(fp)
        doc=PDFDocument(parser)
        parser.set_document(doc)
        resource=PDFResourceManager()
        laparam=LAParams()
        device=PDFPageAggregator(resource,laparams=laparam)
        interpreter=PDFPageInterpreter(resource,device)
        for page in PDFPage.create_pages(doc):
            interpreter.process_page(page)
            layout=device.get_result()
            for out in layout:
                if hasattr(out, "get_text"):
                    para_text += out.get_text()
        return para_text






    def get_text_docx(self,JD_path_name):
        "获取doc内所有字符串"
        doc = docx.Document(JD_path_name)
        para_text = ""
        for paragraph in doc.paragraphs:
            # print(paragraph.text)
            para_text += paragraph.text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    para_text += cell.text
        # print(para_text)
        return para_text







    def get_all_JD_from_lib(self,floder_path, key_world_list):
        self.find_file(floder_path)

        
    def find_file(self,path, file_list=[]):
        """
        加载所有文件
        """
        dir = os.listdir(path)
        for i in dir:
            i = os.path.join(path, i)
            if os.path.isdir(i):
                find_file(i, file_list)
            else:
                file_list.append(i)
                print(file_list)
        return file_list


    
            
    
    



if __name__=="__main__":
    lib_path = "C:\\Users\\李子汉\\Desktop\\intv\\JD_lib_docx"
    key_world_list = ["郭",'性别']
    # load_file().get_JD_msg(lib_path)
    load_file().get_list_from_lib(lib_path,key_world_list)
    